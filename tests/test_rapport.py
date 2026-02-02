#!/usr/bin/env python3
"""
Script de test pour la génération de rapport Word.
Simule le workflow complet sans passer par Streamlit ni les overrides.

Usage:
    python tests/test_generation_rapport.py <fichier_excel.xlsx>

Le rapport généré sera dans : tests/output/rapport_test.docx
"""

import sys
import os
from pathlib import Path

# Ajouter le dossier parent au path pour importer les modules
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document

from src.flat_loader import FlatLoader, ExcelValidationError
from src.tree import OrganizationTree
from src.calc_emissions import EmissionCalculator, EmissionOverrides
from src.calc_indicators import IndicatorCalculator
from src.content_catalog import ContentCatalog
from src.kpi_calculators import KPICalculator
from src.word_renderer import WordRenderer


def _collect_doc_text(doc: Document) -> str:
    """Récupère le texte des paragraphes et des cellules de tableau."""
    chunks = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            chunks.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        chunks.append(paragraph.text)
    return "\n".join(chunks)


def _expected_activity_counts(tree: OrganizationTree) -> dict:
    """Calcule le nombre d'occurrences attendu par activité."""
    if not tree.has_lots():
        return {activity: 1 for activity in tree.get_org_activities()}

    counts = {}
    for lot in tree.get_lots():
        for activity in tree.get_lot_activities(lot.node_id):
            counts[activity] = counts.get(activity, 0) + 1
    return counts


def test_generation_rapport(excel_path: str, output_path: str = None, annee: int = 2024):
    """
    Teste la génération complète d'un rapport.

    Args:
        excel_path: Chemin vers le fichier Excel
        output_path: Chemin de sortie (optionnel)
        annee: Année du bilan

    Returns:
        True si succès, False sinon
    """
    print("=" * 70)
    print("🧪 TEST DE GÉNÉRATION DE RAPPORT CARBONE")
    print("=" * 70)
    print()

    # Définir le chemin de sortie (sera mis à jour après construction de l'arbre)
    if output_path is None:
        output_dir = Path(__file__).parent / "output"
        output_dir.mkdir(exist_ok=True)
        output_path = output_dir / "rapport_test.docx"  # Temporaire, sera mis à jour
    else:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        # 1. CHARGEMENT EXCEL
        print("📥 Étape 1/7 : Chargement du fichier Excel...")
        print(f"   Fichier : {excel_path}")

        loader = FlatLoader(excel_path)
        data = loader.load()

        errors, warnings = loader.get_validation_report()
        if warnings:
            print(f"   ⚠️  {len(warnings)} avertissement(s)")
            for warning in warnings[:3]:  # Afficher max 3
                print(f"      • {warning}")

        print("   ✅ Excel chargé et validé")
        print()

        # 2. CONSTRUCTION ARBORESCENCE
        print("🌳 Étape 2/7 : Construction de l'arborescence...")

        tree = OrganizationTree(data['ORG_TREE'])

        # Validation
        tree_errors = tree.validate_structure()
        if tree_errors:
            print("   ❌ Erreurs dans la structure :")
            for error in tree_errors:
                print(f"      • {error}")
            return False

        org = tree.get_org()
        lots = tree.get_lots()
        ents = tree.get_ents()

        print(f"   ORG : {org.node_name}")
        print(f"   LOTs : {len(lots)}")
        print(f"   ENTs : {len(ents)}")
        print(f"   Activités : {', '.join(sorted(tree.get_org_activities()))}")
        print("   ✅ Arborescence construite")
        print()

        # Mettre à jour le nom du fichier de sortie avec le nom de l'organisation
        if output_path.name == "rapport_test.docx":
            import re
            org_name_clean = re.sub(r'[^\w\s-]', '', org.node_name).strip().replace(' ', '_')
            output_path = output_path.parent / f"Rapport Bilan Carbone {org_name_clean} {annee}.docx"

        # 3. CALCULS ÉMISSIONS
        print("📊 Étape 3/7 : Calcul des émissions...")

        emission_calc = EmissionCalculator(
            tree,
            data['EMISSIONS'],
            data['POSTES_REF']
        )

        # Calcul avec overrides auto (exclusion chauffage AEP si applicable)
        auto_overrides = loader.get_auto_overrides()
        if auto_overrides.poste_config:
            results_brut = emission_calc.calculate_net(auto_overrides, top_n=4)
        else:
            results_brut = emission_calc.calculate_brut(top_n=4)

        org_result = results_brut.get('ORG')
        if org_result:
            print(f"   Total ORG : {org_result.total_tco2e:.1f} tCO₂e")
            print(f"   • Scope 1 : {org_result.scope1_tco2e:.1f} tCO₂e")
            print(f"   • Scope 2 : {org_result.scope2_tco2e:.1f} tCO₂e")
            print(f"   • Scope 3 : {org_result.scope3_tco2e:.1f} tCO₂e")
            if org_result.top_postes:
                print(f"   Top poste : {emission_calc.get_poste_label(org_result.top_postes[0][0])}")

        print("   ✅ Émissions calculées")
        print()

        # 4. CALCULS INDICATEURS
        print("📈 Étape 4/7 : Calcul des indicateurs...")

        indicator_calc = IndicatorCalculator(
            tree,
            data['INDICATORS'],
            data['INDICATORS_REF']
        )

        indicator_results = indicator_calc.calculate()

        print(f"   Indicateurs calculés : {len(indicator_results)} périmètres")
        print("   ✅ Indicateurs calculés")
        print()

        # 5. CHARGEMENT CATALOGUE CONTENUS
        print("📖 Étape 5/7 : Chargement du catalogue de contenus...")

        content_catalog = ContentCatalog(data['TEXTE_RAPPORT'])

        postes_catalog = content_catalog.get_all_postes()
        print(f"   Postes dans le catalogue : {len(postes_catalog)}")
        print("   ✅ Catalogue chargé")
        print()

        # 6. CALCUL DES KPI
        print("🔢 Étape 6/7 : Calcul des KPI...")

        kpi_calc = KPICalculator()

        # KPI globaux
        if org_result:
            kpi_flights = kpi_calc.calculate_flight_equivalent(org_result.total_tco2e)
            kpi_persons = kpi_calc.calculate_person_equivalent(org_result.total_tco2e)
            print(f"   Équivalent vols : {kpi_flights:.0f}")
            print(f"   Équivalent personnes : {kpi_persons:.0f}")

        # KPI m³ par activité - CORRIGÉ
        kpi_m3_eu = None
        kpi_m3_aep = None

        # Chercher les résultats EU et AEP
        eu_result = None
        aep_result = None
        eu_indicators_list = []
        aep_indicators_list = []

        for key, result in results_brut.items():
            if result.activity == 'EU':
                eu_result = result
                indicator_res = indicator_results.get(key)
                if indicator_res:
                    eu_indicators_list.append(indicator_res)
            elif result.activity == 'AEP':
                aep_result = result
                indicator_res = indicator_results.get(key)
                if indicator_res:
                    aep_indicators_list.append(indicator_res)

        # Calculer KPI m³ en sommant TOUS les volumes
        if eu_result and eu_indicators_list:
            kpi_m3_eu = kpi_calc.calculate_kpi_m3_eu(eu_result, eu_indicators_list)
            if kpi_m3_eu:
                print(f"   KPI EU : {kpi_m3_eu:.2f} kgCO₂e/m³")

        if aep_result and aep_indicators_list:
            kpi_m3_aep = kpi_calc.calculate_kpi_m3_aep(aep_result, aep_indicators_list)
            if kpi_m3_aep:
                print(f"   KPI AEP : {kpi_m3_aep:.2f} kgCO₂e/m³")

        # Texte de comparaison (utiliser le premier indicateur de chaque liste)
        eu_ind_first = eu_indicators_list[0] if eu_indicators_list else None
        aep_ind_first = aep_indicators_list[0] if aep_indicators_list else None
        activity_comparison = kpi_calc.generate_activity_volume_comparison_text(
            eu_result, aep_result, eu_ind_first, aep_ind_first
        )

        print("   ✅ KPI calculés")
        print()

        # 7. GÉNÉRATION DU RAPPORT WORD
        print("📝 Étape 7/7 : Génération du rapport Word...")

        template_path = Path("templates/rapport_template.docx")
        if not template_path.exists():
            print(f"   ❌ Template non trouvé : {template_path}")
            print(f"   → Placer votre template dans templates/rapport_template.docx")
            return False

        # Calculer les données chauffage AEP
        aep_with_chauffage = emission_calc.calculate_aep_with_chauffage()
        chauffage_total = emission_calc.get_chauffage_total()
        org_with_chauffage = emission_calc.calculate_org_with_chauffage()

        # Préparer le contexte
        context = {
            'annee': annee,
            'org_result': org_result,
            'lot_results': {k: v for k, v in results_brut.items() if k.startswith('LOT_') or k.startswith('ORG_')},
            'has_lots': tree.has_lots(),
            'poste_labels': emission_calc.poste_labels,
            'top_n': 4,
            'overrides': EmissionOverrides(),  # Pas d'overrides pour le test
            'kpi_m3_eu': kpi_m3_eu,  # CORRIGÉ : nouveaux noms
            'kpi_m3_aep': kpi_m3_aep,  # CORRIGÉ : nouveaux noms
            'activity_volume_comparison_text': activity_comparison,
            'indicator_results': indicator_results,
            'content_catalog': content_catalog,
            'emissions_l2_df': data.get('EMISSIONS_L2'),
            'emissions_df': data.get('EMISSIONS'),
            'tree': tree,
            'aep_with_chauffage_result': aep_with_chauffage,
            'chauffage_total_tco2e': chauffage_total,
            'org_with_chauffage_result': org_with_chauffage,
            'beges_df': data.get('BEGES'),
            'emissions_evitees_df': data.get('EMISSIONS_EVITEES'),
        }

        # Générer le rapport
        renderer = WordRenderer(
            template_path=str(template_path),
            assets_path="assets"
        )

        print(f"   Template : {template_path}")
        doc = renderer.render(context)

        # Sauvegarder
        renderer.save(str(output_path))

        print(f"   ✅ Rapport généré : {output_path}")
        print()

        # Vérifier la répétition LOT/ACTIVITY
        doc_check = Document(str(output_path))
        full_text = _collect_doc_text(doc_check)

        assert '[[START_' not in full_text, "Marqueurs START encore présents dans le rapport"
        assert '[[END_' not in full_text, "Marqueurs END encore présents dans le rapport"

        lot_names = [lot.node_name for lot in lots]
        missing_lots = [name for name in lot_names if name not in full_text]
        assert not missing_lots, f"Lots manquants dans le rapport: {missing_lots}"

        activity_labels = {"EU": "Eaux usées", "AEP": "Eau potable"}
        expected_counts = _expected_activity_counts(tree)
        for activity, expected in expected_counts.items():
            label = activity_labels.get(activity)
            if label:
                needle = f"Activité – {label}"
                found = full_text.count(needle)
                assert found == expected, f"Activité {activity} attendue {expected}, trouvé {found}"

        # RÉSUMÉ
        print("=" * 70)
        print("✅ GÉNÉRATION RÉUSSIE !")
        print("=" * 70)
        print()
        print(f"📄 Rapport généré : {output_path.absolute()}")
        print(f"📊 Émissions totales : {org_result.total_tco2e:.1f} tCO₂e")
        print(f"🌳 Structure : {len(lots)} LOT(s), {len(ents)} ENT(s)")
        print(f"📈 Activités : {', '.join(sorted(tree.get_org_activities()))}")
        print()

        return True

    except ExcelValidationError as e:
        print(f"❌ Erreur de validation Excel :")
        print(f"   {str(e)}")
        return False

    except Exception as e:
        print(f"❌ Erreur lors de la génération :")
        print(f"   {str(e)}")
        import traceback
        print()
        print("📋 Traceback complet :")
        traceback.print_exc()
        return False


def main():
    """Fonction principale."""
    if len(sys.argv) < 2:
        print("Usage: python tests/test_generation_rapport.py <fichier_excel.xlsx> [annee]")
        print()
        print("Exemple:")
        print("  python tests/test_generation_rapport.py data/mon_bilan.xlsx 2024")
        print()
        sys.exit(1)

    excel_path = sys.argv[1]
    annee = int(sys.argv[2]) if len(sys.argv) > 2 else 2024

    if not Path(excel_path).exists():
        print(f"❌ Fichier non trouvé : {excel_path}")
        sys.exit(1)

    success = test_generation_rapport(excel_path, annee=annee)

    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
