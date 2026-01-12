#!/usr/bin/env python3
"""
Tests unitaires pour les nouvelles fonctionnalités de word_renderer.
Tests de l'insertion du logo et des blocs répétables.
"""

import sys
import tempfile
from pathlib import Path

import pandas as pd
from docx import Document

# Ajouter le dossier racine au path (2 niveaux au-dessus car on est dans tests/unit/)
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from src.word_renderer import WordRenderer
from src.tree import OrganizationTree
from src.calc_emissions import EmissionResult


def _build_test_tree() -> OrganizationTree:
    """Construit une arborescence simple ORG -> LOT -> ENT avec EU/AEP."""
    data = [
        {"node_id": "ORG1", "parent_id": None, "node_type": "ORG", "node_name": "ORG", "activity": None},
        {"node_id": "LOT1", "parent_id": "ORG1", "node_type": "LOT", "node_name": "Lot A", "activity": None},
        {"node_id": "LOT2", "parent_id": "ORG1", "node_type": "LOT", "node_name": "Lot B", "activity": None},
        {"node_id": "ENT1", "parent_id": "LOT1", "node_type": "ENT", "node_name": "Ent 1", "activity": "EU"},
        {"node_id": "ENT2", "parent_id": "LOT1", "node_type": "ENT", "node_name": "Ent 2", "activity": "AEP"},
        {"node_id": "ENT3", "parent_id": "LOT2", "node_type": "ENT", "node_name": "Ent 3", "activity": "EU"},
        {"node_id": "ENT4", "parent_id": "LOT2", "node_type": "ENT", "node_name": "Ent 4", "activity": "AEP"},
    ]
    return OrganizationTree(pd.DataFrame(data))


def _build_minimal_template(path: Path) -> None:
    """Cree un template minimal avec blocs LOT/ACTIVITY."""
    doc = Document()
    doc.add_paragraph("[[START_LOT]]")
    doc.add_paragraph("{{LOT_NAME}}")
    doc.add_paragraph("[[START_ACTIVITY]]")
    doc.add_paragraph("Activité – {{ENT_ACTIVITY}}")
    doc.add_paragraph("[[END_ACTIVITY]]")
    doc.add_paragraph("[[END_LOT]]")
    doc.save(path)


def _collect_doc_text(doc: Document) -> str:
    """Recupere le texte des paragraphes et des cellules de tableau."""
    chunks = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            chunks.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        chunks.append(paragraph.text)
    return "\n".join(chunks)


def test_word_renderer_initialization():
    """Teste que WordRenderer s'initialise correctement avec les nouvelles méthodes."""
    template_path = "templates/rapport_template.docx"
    assets_path = "assets"

    renderer = WordRenderer(template_path, assets_path)

    # Vérifier que toutes les nouvelles méthodes existent
    assert hasattr(renderer, '_insert_static_logo'), "Méthode _insert_static_logo manquante"
    assert hasattr(renderer, '_insert_asset_image'), "Méthode _insert_asset_image manquante"
    assert hasattr(renderer, '_process_lot_blocks'), "Méthode _process_lot_blocks manquante"
    assert hasattr(renderer, '_process_org_activity_blocks'), "Méthode _process_org_activity_blocks manquante"
    assert hasattr(renderer, '_process_activity_blocks'), "Méthode _process_activity_blocks manquante"
    assert hasattr(renderer, '_process_post_blocks'), "Méthode _process_post_blocks manquante"
    assert hasattr(renderer, '_insert_post_content'), "Méthode _insert_post_content manquante"
    assert hasattr(renderer, '_find_all_lot_blocks'), "Méthode _find_all_lot_blocks manquante"
    assert hasattr(renderer, '_find_all_activity_blocks'), "Méthode _find_all_activity_blocks manquante"
    assert hasattr(renderer, '_find_all_post_blocks'), "Méthode _find_all_post_blocks manquante"
    assert hasattr(renderer, '_find_all_other_post_blocks'), "Méthode _find_all_other_post_blocks manquante"
    assert hasattr(renderer, '_process_other_post_blocks'), "Méthode _process_other_post_blocks manquante"
    assert hasattr(renderer, '_insert_entity_charts'), "Méthode _insert_entity_charts manquante"
    assert hasattr(renderer, '_delete_block'), "Méthode _delete_block manquante"
    assert hasattr(renderer, '_clean_all_markers'), "Méthode _clean_all_markers manquante"
    assert hasattr(renderer, '_generate_post_chart'), "Méthode _generate_post_chart manquante"
    assert hasattr(renderer, '_insert_post_table'), "Méthode _insert_post_table manquante"

    print("✅ Toutes les méthodes sont présentes dans WordRenderer")
    return True


def test_logo_insertion_method_exists():
    """Teste que la méthode d'insertion du logo existe."""
    template_path = "templates/rapport_template.docx"
    assets_path = "assets"

    renderer = WordRenderer(template_path, assets_path)

    # Vérifier que la méthode _insert_static_logo existe
    assert callable(renderer._insert_static_logo), "_insert_static_logo n'est pas callable"

    print("✅ Méthode _insert_static_logo est callable")
    return True


def test_block_processing_methods_exist():
    """Teste que les méthodes de traitement des blocs existent."""
    template_path = "templates/rapport_template.docx"
    assets_path = "assets"

    renderer = WordRenderer(template_path, assets_path)

    # Vérifier que toutes les méthodes de traitement des blocs sont callable
    assert callable(renderer._process_lot_blocks), "_process_lot_blocks n'est pas callable"
    assert callable(renderer._process_activity_blocks), "_process_activity_blocks n'est pas callable"
    assert callable(renderer._process_post_blocks), "_process_post_blocks n'est pas callable"
    assert callable(renderer._clean_all_markers), "_clean_all_markers n'est pas callable"

    print("✅ Toutes les méthodes de traitement des blocs sont callable")
    return True


def test_repetition_lot_activity_blocks():
    """Teste que la repetition LOT/ACTIVITY fonctionne sur un template minimal."""
    with tempfile.TemporaryDirectory() as tmp_dir:
        template_path = Path(tmp_dir) / "template_test.docx"
        output_path = Path(tmp_dir) / "output_test.docx"
        _build_minimal_template(template_path)

        tree = _build_test_tree()
        lot_results = {}
        for lot in tree.get_lots():
            for activity in ["EU", "AEP"]:
                key = f"LOT_{lot.node_id}_{activity}"
                lot_results[key] = EmissionResult(
                    node_id=lot.node_id,
                    node_name=lot.node_name,
                    activity=activity,
                )

        renderer = WordRenderer(str(template_path), "assets")
        context = {
            "annee": 2024,
            "org_result": None,
            "lot_results": lot_results,
            "has_lots": True,
            "poste_labels": {},
            "top_n": 4,
            "overrides": None,
            "kpi_m3_eu": None,
            "kpi_m3_aep": None,
            "activity_volume_comparison_text": "",
            "indicator_results": {},
            "content_catalog": None,
            "emissions_l2_df": None,
            "tree": tree,
        }

        renderer.render(context)
        renderer.save(str(output_path))

        doc = Document(str(output_path))
        text = _collect_doc_text(doc)

        assert "[[START_" not in text and "[[END_" not in text, "Marqueurs encore presents"
        assert "Lot A" in text and "Lot B" in text, "Noms de LOT manquants"
        assert text.count("Activité – Eau potable") == 2, "Nombre d'activites AEP incorrect"
        assert text.count("Activité – Eaux usées") == 2, "Nombre d'activites EU incorrect"

        print("✅ Repetition LOT/ACTIVITY OK")
        return True


def main():
    """Exécute tous les tests."""
    print("=" * 70)
    print("🧪 TESTS UNITAIRES - WordRenderer - Blocs répétables")
    print("=" * 70)
    print()

    tests = [
        ("Initialisation de WordRenderer", test_word_renderer_initialization),
        ("Méthode d'insertion du logo", test_logo_insertion_method_exists),
        ("Méthodes de traitement des blocs", test_block_processing_methods_exist),
        ("Répétition LOT/ACTIVITY", test_repetition_lot_activity_blocks),
    ]

    passed = 0
    failed = 0

    for test_name, test_func in tests:
        print(f"🔍 Test: {test_name}")
        try:
            result = test_func()
            if result:
                passed += 1
                print()
            else:
                failed += 1
                print(f"❌ Test échoué: {test_name}")
                print()
        except Exception as e:
            failed += 1
            print(f"❌ Test échoué avec erreur: {test_name}")
            print(f"   Erreur: {str(e)}")
            print()

    print("=" * 70)
    print(f"📊 Résultats: {passed} réussis, {failed} échoués")
    print("=" * 70)

    return 0 if failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
