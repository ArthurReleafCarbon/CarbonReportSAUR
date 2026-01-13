"""
Module helper pour gérer les blocs répétables dans Word.
Gère la duplication de blocs [[START_XXX]] ... [[END_XXX]].
"""

import re
from docx.document import Document
from copy import deepcopy
from typing import List, Tuple, Optional


class BlockProcessor:
    """Processeur de blocs répétables dans un document Word."""

    def __init__(self, doc: Document):
        """
        Initialise le processeur.

        Args:
            doc: Document Word
        """
        self.doc = doc

    def find_block(self, start_marker: str, end_marker: str) -> Optional[Tuple[int, int]]:
        """
        Trouve un bloc dans le document.

        Args:
            start_marker: Marqueur de début (ex: '[[START_LOT]]')
            end_marker: Marqueur de fin (ex: '[[END_LOT]]')

        Returns:
            Tuple (index_start, index_end) des paragraphes ou None si non trouvé
        """
        start_idx = None
        end_idx = None

        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text.strip()

            if start_marker in text:
                start_idx = i
            elif end_marker in text and start_idx is not None:
                end_idx = i
                break

        if start_idx is not None and end_idx is not None:
            return (start_idx, end_idx)

        return None

    def extract_block_content(self, start_idx: int, end_idx: int) -> List:
        """
        Extrait le contenu d'un bloc (paragraphes entre start et end, exclus).

        Args:
            start_idx: Index du paragraphe START
            end_idx: Index du paragraphe END

        Returns:
            Liste des éléments du bloc
        """
        # Récupérer les éléments XML entre start et end
        elements = []
        paragraphs = list(self.doc.paragraphs)

        for i in range(start_idx + 1, end_idx):
            if i < len(paragraphs):
                elements.append(paragraphs[i]._element)

        return elements

    def duplicate_block(self, start_idx: int, end_idx: int, n_copies: int) -> List[Tuple[int, int]]:
        """
        Duplique un bloc n fois AVEC ses marqueurs START et END.

        Args:
            start_idx: Index du paragraphe START
            end_idx: Index du paragraphe END
            n_copies: Nombre de copies à créer

        Returns:
            Liste des tuples (start_idx, end_idx) pour chaque copie
        """
        paragraphs = list(self.doc.paragraphs)

        # Extraire TOUS les éléments du bloc (START inclus, END inclus)
        block_elements = []
        for i in range(start_idx, end_idx + 1):
            if i < len(paragraphs):
                block_elements.append(paragraphs[i]._element)

        # Position d'insertion (après END_MARKER)
        insert_after = paragraphs[end_idx]._element

        # Dupliquer n fois
        copies_indices = []
        current_insert_after = insert_after

        for copy_num in range(n_copies):
            # Copier chaque élément (y compris START et END)
            for element in block_elements:
                # Cloner l'élément
                new_element = deepcopy(element)

                # Insérer après la position courante
                current_insert_after.addnext(new_element)
                current_insert_after = new_element

        return copies_indices

    def remove_block_markers(self, start_marker: str, end_marker: str):
        """
        Supprime les marqueurs de bloc du document.

        Args:
            start_marker: Marqueur de début
            end_marker: Marqueur de fin
        """
        paragraphs_to_remove = []

        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if text == start_marker or text == end_marker:
                paragraphs_to_remove.append(paragraph)

        # Supprimer les paragraphes
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)

    def replace_in_block(self, start_idx: int, end_idx: int, replacements: dict):
        """
        Remplace des placeholders dans un bloc (paragraphes ET tableaux).

        Args:
            start_idx: Index de début du bloc (basé sur les paragraphes)
            end_idx: Index de fin du bloc (basé sur les paragraphes)
            replacements: Dictionnaire {placeholder: valeur}
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        paragraphs = list(self.doc.paragraphs)

        # 1. Remplacer dans les paragraphes
        for i in range(start_idx, end_idx + 1):
            if i < len(paragraphs):
                paragraph = paragraphs[i]
                text = paragraph.text

                # Remplacer les placeholders
                for placeholder, value in replacements.items():
                    text = text.replace(placeholder, str(value))

                    # Forcer l'alignement à gauche pour les listes formatées
                    if placeholder == '{{ENTITY_TOP_POSTES_LIST}}' and placeholder in paragraph.text:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Mettre à jour le texte si modifié
                if text != paragraph.text:
                    # Conserver le formatage
                    if paragraph.runs:
                        paragraph.runs[0].text = text
                        for run in paragraph.runs[1:]:
                            run.text = ''
                    else:
                        paragraph.text = text

        # 2. Remplacer dans les tableaux qui se trouvent dans le bloc
        # Trouver les tableaux entre start_idx et end_idx
        if start_idx >= len(paragraphs) or end_idx >= len(paragraphs):
            return

        start_element = paragraphs[start_idx]._element
        end_element = paragraphs[end_idx]._element

        # Parcourir tous les tableaux du document
        for table in self.doc.tables:
            table_element = table._element

            # Vérifier si le tableau est dans la plage du bloc
            # (entre start_element et end_element dans le body)
            is_in_range = self._is_element_in_range(table_element, start_element, end_element)

            if is_in_range:
                # Remplacer dans toutes les cellules du tableau
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text

                            # Remplacer les placeholders
                            for placeholder, value in replacements.items():
                                text = text.replace(placeholder, str(value))

                            # Mettre à jour le texte si modifié
                            if text != paragraph.text:
                                if paragraph.runs:
                                    paragraph.runs[0].text = text
                                    for run in paragraph.runs[1:]:
                                        run.text = ''
                                else:
                                    paragraph.text = text

    def _is_element_in_range(self, element, start_element, end_element):
        """
        Vérifie si un élément XML est entre start_element et end_element dans le body.

        Args:
            element: Élément à vérifier
            start_element: Élément de début
            end_element: Élément de fin

        Returns:
            True si l'élément est dans la plage, False sinon
        """
        body = self.doc.element.body
        elements = list(body)

        try:
            start_pos = elements.index(start_element)
            end_pos = elements.index(end_element)
            element_pos = elements.index(element)

            return start_pos <= element_pos <= end_pos
        except ValueError:
            # Un des éléments n'est pas dans le body
            return False


class BlockTemplate:
    """Template pour un bloc répétable."""

    def __init__(self, start_marker: str, end_marker: str):
        """
        Initialise le template.

        Args:
            start_marker: Marqueur de début
            end_marker: Marqueur de fin
        """
        self.start_marker = start_marker
        self.end_marker = end_marker
        self.instances = []

    def add_instance(self, replacements: dict):
        """
        Ajoute une instance du bloc avec ses remplacements.

        Args:
            replacements: Dictionnaire de remplacements pour cette instance
        """
        self.instances.append(replacements)
