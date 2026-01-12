"""
Module de rendu Word.
Gère le remplacement des placeholders, la duplication des blocs répétables,
l'insertion d'images et tableaux, et le nettoyage des placeholders vides.
"""

import re
from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path
from io import BytesIO

from .calc_emissions import EmissionResult, EmissionOverrides
from .calc_indicators import IndicatorResult
from .chart_generators import ChartGenerator
from .table_generators import TableGenerator
from .kpi_calculators import KPICalculator
from .content_catalog import ContentCatalog


class WordRenderer:
    """
    Moteur de rendu Word.
    Remplace les placeholders, duplique les blocs, insère images/tableaux.
    """

    def __init__(self, template_path: str, assets_path: str):
        """
        Initialise le renderer.

        Args:
            template_path: Chemin vers le template Word
            assets_path: Chemin vers le dossier assets
        """
        self.template_path = Path(template_path)
        self.assets_path = Path(assets_path)
        self.doc = None

        # Générateurs
        self.chart_gen = ChartGenerator()
        self.table_gen = TableGenerator()
        self.kpi_calc = KPICalculator()

    def load_template(self):
        """Charge le template Word."""
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template non trouvé : {self.template_path}")
        self.doc = Document(self.template_path)

    def render(self, context: Dict[str, Any]) -> Document:
        """
        Effectue le rendu complet du document.

        Args:
            context: Dictionnaire de contexte contenant toutes les données

        Returns:
            Document Word rendu
        """
        self.load_template()

        # 1. Remplacer les placeholders simples globaux
        self._replace_simple_placeholders(context)

        # 1.5 Insérer le logo statique
        self._insert_static_logo()

        # 2. Dupliquer et remplir les blocs LOT
        self._process_lot_blocks(context)

        # 2.5 Nettoyer tous les marqueurs de blocs
        self._clean_all_markers()

        # 3. Insérer les graphiques ORG
        self._insert_org_charts(context)

        # 4. Nettoyer les placeholders vides
        self._clean_empty_placeholders()

        return self.doc

    def _replace_simple_placeholders(self, context: Dict[str, Any]):
        """
        Remplace les placeholders simples de type {{PLACEHOLDER}}.

        Args:
            context: Dictionnaire de contexte
        """
        replacements = self._build_simple_replacements(context)

        # Parcourir tous les paragraphes
        for paragraph in self.doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

        # Parcourir les tableaux
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)

    def _build_simple_replacements(self, context: Dict[str, Any]) -> Dict[str, str]:
        """
        Construit le dictionnaire des remplacements simples.

        Args:
            context: Contexte global

        Returns:
            Dictionnaire {placeholder: valeur}
        """
        replacements = {}

        # Année
        replacements['{{annee}}'] = str(context.get('annee', '2024'))

        # ORG - AVEC FORMAT ESPACE ET ARRONDI
        org_result = context.get('org_result')
        if org_result:
            replacements['{{ORG_NAME}}'] = org_result.node_name
            replacements['{{TOTAL_EMISSIONS}}'] = self.kpi_calc.format_number(org_result.total_tco2e)
            replacements['{{TOTAL_EMISSIONS_S1}}'] = self.kpi_calc.format_number(org_result.scope1_tco2e)
            replacements['{{TOTAL_EMISSIONS_S2}}'] = self.kpi_calc.format_number(org_result.scope2_tco2e)
            replacements['{{TOTAL_EMISSIONS_S3}}'] = self.kpi_calc.format_number(org_result.scope3_tco2e)
            replacements['{{pourc_s3_org}}'] = self.kpi_calc.format_number(org_result.get_scope_percentage(3))

        # KPI m³ - NOUVEAUX NOMS
        kpi_m3_eu = context.get('kpi_m3_eu')
        kpi_m3_aep = context.get('kpi_m3_aep')
        if kpi_m3_eu is not None:
            replacements['{{kpi_M3_EU}}'] = f"{kpi_m3_eu:.2f} kgCO₂e/m³".replace(".", ",")
        else:
            replacements['{{kpi_M3_EU}}'] = "N/A"

        if kpi_m3_aep is not None:
            replacements['{{kpi_M3_AEP}}'] = f"{kpi_m3_aep:.2f} kgCO₂e/m³".replace(".", ",")
        else:
            replacements['{{kpi_M3_AEP}}'] = "N/A"

        # Équivalents - AVEC FORMAT ESPACE
        if org_result:
            flights = self.kpi_calc.calculate_flight_equivalent(org_result.total_tco2e)
            persons = self.kpi_calc.calculate_person_equivalent(org_result.total_tco2e)
            replacements['{{kpi_1}}'] = self.kpi_calc.format_number(flights)
            replacements['{{kpi_2}}'] = self.kpi_calc.format_number(persons)

        # Top postes ORG
        if org_result and org_result.top_postes:
            top_3 = org_result.top_postes[:3]
            poste_labels = context.get('poste_labels', {})
            for i, (code, _) in enumerate(top_3, 1):
                label = poste_labels.get(code, code)
                replacements[f'{{{{TOP_POSTE_{i}}}}}'] = label

        # Texte de comparaison volumes
        comp_text = context.get('activity_volume_comparison_text', '')
        replacements['{{ACTIVITY_VOLUME_COMPARISON_TEXT}}'] = comp_text

        # Top postes longueur
        top_n = context.get('top_n', 4)
        replacements['{{TOP_POSTES_LONGUEUR}}'] = str(top_n)

        # Comparaison volumes par entité (LOT/activité)
        entity_count, entities = self._get_entity_volume_comparison(context)
        replacements['{{ENTITY_COUNT}}'] = str(entity_count)

        # Générer le texte de comparaison dynamique
        replacements['{{ENTITY_VOLUME_COMPARISON}}'] = self._generate_entity_volume_text(entities)

        # Garder les placeholders individuels pour compatibilité (2 premières entités)
        if len(entities) >= 1:
            replacements['{{ENTITY_1_NAME}}'] = entities[0]['name']
            replacements['{{ENTITY_1_VOLUME_M3}}'] = entities[0]['volume']
        else:
            replacements['{{ENTITY_1_NAME}}'] = ''
            replacements['{{ENTITY_1_VOLUME_M3}}'] = ''
        if len(entities) >= 2:
            replacements['{{ENTITY_2_NAME}}'] = entities[1]['name']
            replacements['{{ENTITY_2_VOLUME_M3}}'] = entities[1]['volume']
        else:
            replacements['{{ENTITY_2_NAME}}'] = ''
            replacements['{{ENTITY_2_VOLUME_M3}}'] = ''

        return replacements

    def _format_volume_millions(self, value: float) -> str:
        """Formate un volume en millions avec virgule française."""
        if value is None:
            return ''
        millions = value / 1_000_000
        rounded = round(millions, 1)
        if abs(rounded - int(rounded)) < 1e-6:
            return f"{int(rounded)} million"
        return f"{str(rounded).replace('.', ',')} million"

    def _format_tco2_value(self, value: float) -> str:
        """Formate une valeur tCO2 avec 2 décimales et virgule française."""
        formatted = self.kpi_calc.format_number(value, decimals=2)
        return formatted.replace(".", ",")

    def _format_entity_top_postes_list(self, result: EmissionResult, poste_labels: Dict[str, str]) -> str:
        """Construit la liste des 4 postes les plus émissifs pour une entité."""
        if not result or not result.top_postes:
            return ''

        lines = []
        for poste_code, tco2e in result.top_postes[:4]:
            label = poste_labels.get(poste_code, poste_code)
            value_str = self._format_tco2_value(tco2e)
            lines.append(f"- {label} - {value_str} t CO2")

        return "\n".join(lines)

    def _generate_entity_volume_text(self, entities: List[Dict]) -> str:
        """
        Génère un texte de comparaison adapté au nombre d'entités.

        Args:
            entities: Liste d'entités triées par volume décroissant

        Returns:
            Texte formaté décrivant la comparaison des volumes
        """
        if not entities:
            return ""

        if len(entities) == 1:
            # Une seule entité
            return (f"L'entité {entities[0]['name']} présente un volume "
                   f"de {entities[0]['volume']} m³.")

        elif len(entities) == 2:
            # Deux entités : comparaison simple
            return (f"À ce titre, {entities[0]['name']} présente un volume "
                   f"de {entities[0]['volume']} m³ d'eau traités contre "
                   f"{entities[1]['volume']} m³ pour {entities[1]['name']}, "
                   f"ce qui influence directement le niveau d'émissions observé "
                   f"pour les postes les plus contributifs.")

        else:
            # Plus de 2 entités : lister les 2 premières + mentionner les autres
            text_parts = []

            # Lister les 2 premières
            text_parts.append(
                f"Les principales entités en termes de volumes sont {entities[0]['name']} "
                f"avec {entities[0]['volume']} m³ et {entities[1]['name']} "
                f"avec {entities[1]['volume']} m³."
            )

            # Mentionner les autres
            if len(entities) == 3:
                text_parts.append(
                    f"L'entité {entities[2]['name']} représente {entities[2]['volume']} m³."
                )
            elif len(entities) > 3:
                others = entities[2:]
                other_names = ", ".join([e['name'] for e in others[:-1]])
                last_name = others[-1]['name']
                text_parts.append(
                    f"Les autres entités ({other_names} et {last_name}) "
                    f"complètent le périmètre avec des volumes plus modestes."
                )

            text_parts.append(
                "Ces différences de volumes influencent directement le niveau "
                "d'émissions observé pour les postes les plus contributifs."
            )

            return " ".join(text_parts)

    def _get_entity_volume_comparison(self, context: Dict[str, Any]):
        """
        Construit les valeurs de comparaison volumes pour les placeholders ENTITY_*.
        Retourne le nombre total d'entités (LOT×ACTIVITÉ) et toutes les entités triées par volume.
        """
        tree = context.get('tree')
        indicator_results = context.get('indicator_results', {})
        if not tree or not indicator_results:
            return 0, []

        entities = []

        for result in indicator_results.values():
            indicator = None
            if result.activity == 'EU':
                indicator = result.get_indicator('VOL_EAU_EPURE')
            elif result.activity == 'AEP':
                indicator = result.get_indicator('VOL_EAU_DISTRIB')

            if indicator is None:
                # Fallback : chercher n'importe quel indicateur de volume
                for ind in result.indicators.values():
                    unit = (ind.unit or '').lower()
                    label = (ind.indicator_label or '').lower()
                    if 'm3' in unit or 'm³' in unit or 'volume' in label:
                        indicator = ind
                        break

            if indicator is None:
                continue

            entities.append({
                'name': f"{result.node_name} ({result.activity})",
                'volume': self._format_volume_millions(indicator.value),
                'value': indicator.value
            })

        # Trier par volume décroissant
        entities.sort(key=lambda x: x['value'], reverse=True)

        # Le compte d'entités = nombre total d'entités trouvées avec un volume
        entity_count = len(entities)

        # Retourner le compte total et TOUTES les entités (pour texte dynamique)
        return entity_count, entities

    def _replace_in_paragraph(self, paragraph, replacements: Dict[str, str]):
        """
        Remplace les placeholders dans un paragraphe.

        Args:
            paragraph: Paragraphe Word
            replacements: Dictionnaire de remplacements
        """
        # Utiliser le texte complet du paragraphe pour gérer les placeholders fragmentés
        full_text = paragraph.text

        # Vérifier s'il y a des placeholders
        if '{{' not in full_text:
            return

        # Remplacer les placeholders
        new_text = full_text
        for placeholder, value in replacements.items():
            new_text = new_text.replace(placeholder, value)

        # Si le texte a changé, le mettre à jour
        if new_text != full_text:
            # Conserver le style du premier run
            if paragraph.runs:
                style = paragraph.runs[0].style
                font = paragraph.runs[0].font
                # Supprimer tous les runs
                for run in paragraph.runs:
                    run.text = ''
                # Ajouter le nouveau texte
                new_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                new_run.text = new_text
                if style:
                    new_run.style = style
            else:
                paragraph.text = new_text

    def _find_all_lot_blocks(self) -> List[Tuple[int, int]]:
        """
        Trouve tous les blocs LOT dans le document.

        Returns:
            Liste de tuples (start_idx, end_idx) pour chaque bloc LOT
        """
        blocks = []
        current_idx = 0
        paragraphs = list(self.doc.paragraphs)

        while current_idx < len(paragraphs):
            start_idx = None
            end_idx = None

            for i in range(current_idx, len(paragraphs)):
                text = paragraphs[i].text.strip()

                if '[[START_LOT]]' in text and start_idx is None:
                    start_idx = i
                elif '[[END_LOT]]' in text and start_idx is not None:
                    end_idx = i
                    break

            if start_idx is not None and end_idx is not None:
                blocks.append((start_idx, end_idx))
                current_idx = end_idx + 1
            else:
                break

        return blocks

    def _process_lot_blocks(self, context: Dict[str, Any]):
        """
        Traite les blocs LOT répétables.

        Args:
            context: Contexte global
        """
        from .word_blocks import BlockProcessor

        has_lots = context.get('has_lots', False)
        tree = context.get('tree')

        if not has_lots:
            # Cas sans LOT : traiter directement les blocs ACTIVITY au niveau ORG
            self._process_org_activity_blocks(context)
            return

        # Trouver le bloc LOT
        processor = BlockProcessor(self.doc)
        block_info = processor.find_block('[[START_LOT]]', '[[END_LOT]]')

        if not block_info:
            # Pas de bloc LOT dans le template
            return

        start_idx, end_idx = block_info

        # Récupérer les LOTs
        lots = tree.get_lots() if tree else []

        if not lots:
            # Pas de LOTs, supprimer le bloc
            self._delete_block(start_idx, end_idx)
            return

        # Dupliquer le bloc pour chaque LOT (sauf le premier)
        if len(lots) > 1:
            processor.duplicate_block(start_idx, end_idx, len(lots) - 1)

        # Re-scanner pour obtenir tous les blocs LOT
        all_lot_blocks = self._find_all_lot_blocks()

        # Remplir chaque bloc LOT (itération inverse pour éviter les décalages d'index)
        for i in range(len(lots) - 1, -1, -1):
            if i >= len(all_lot_blocks):
                continue

            lot = lots[i]
            block_start, block_end = all_lot_blocks[i]

            # Remplacer {{LOT_NAME}}
            replacements = {
                '{{LOT_NAME}}': lot.node_name
            }

            processor.replace_in_block(block_start, block_end, replacements)

            # Traiter les blocs ACTIVITY imbriqués
            self._process_activity_blocks(block_start, block_end, lot.node_id, context)

    def _process_org_activity_blocks(self, context: Dict[str, Any]):
        """
        Traite les blocs ACTIVITY au niveau ORG (cas sans LOT).

        Args:
            context: Contexte global
        """
        from .word_blocks import BlockProcessor

        # Dans ce cas, les blocs ACTIVITY sont au niveau racine
        # On cherche les blocs ACTIVITY dans tout le document
        processor = BlockProcessor(self.doc)
        block_info = processor.find_block('[[START_ACTIVITY]]', '[[END_ACTIVITY]]')

        if not block_info:
            return

        # Traiter comme si c'était un bloc LOT virtuel
        # Zone = tout le document
        tree = context.get('tree')
        parent_tree_id = tree.get_org().node_id if tree else None
        paragraphs = self.doc.paragraphs
        self._process_activity_blocks(0, len(paragraphs), 'ORG', context, parent_tree_id=parent_tree_id)

    def _insert_asset_image(self, placeholder: str, image_key: str,
                            width: Optional[float] = 5.0, height: Optional[float] = None):
        """
        Insère une image depuis assets/ à la place d'un placeholder.

        Args:
            placeholder: Placeholder à remplacer (ex: '{{POST_IMAGE_1}}')
            image_key: Nom du fichier image (ex: 'DIGESTEUR_SCHEMA')
            width: Largeur en inches
        """
        image_path = self.assets_path / f"{image_key}.png"

        if not image_path.exists():
            # Image non trouvée, laisser le placeholder (sera nettoyé)
            return

        # Charger l'image
        with open(image_path, 'rb') as f:
            img_bytes = f.read()

        img_buffer = BytesIO(img_bytes)

        # Insérer l'image
        self._insert_image(placeholder, img_buffer, width=width, height=height)

    def _insert_static_logo(self):
        """Insère le logo statique ORG_LOGO.png."""
        # Hauteur max 4,5 cm -> 1.77 inches, conserver le ratio
        self._insert_asset_image('{{ORG_LOGO}}', 'ORG_LOGO', width=None, height=1.77)

    def _find_all_post_blocks(self, parent_start: int, parent_end: int) -> List[Tuple[int, int]]:
        """
        Trouve tous les blocs POST dans une zone donnée.

        Args:
            parent_start: Index de début de la zone
            parent_end: Index de fin de la zone

        Returns:
            Liste de tuples (start_idx, end_idx) pour chaque bloc POST
        """
        blocks = []
        current_idx = parent_start
        paragraphs = list(self.doc.paragraphs)

        while current_idx < parent_end and current_idx < len(paragraphs):
            start_idx = None
            end_idx = None

            for i in range(current_idx, min(parent_end, len(paragraphs))):
                text = paragraphs[i].text.strip()

                if '[[START_POST]]' in text and start_idx is None:
                    start_idx = i
                elif '[[END_POST]]' in text and start_idx is not None:
                    end_idx = i
                    break

            if start_idx is not None and end_idx is not None:
                blocks.append((start_idx, end_idx))
                current_idx = end_idx + 1
            else:
                break

        return blocks

    def _resolve_post_content(self, poste_code: str, poste_label: str, activity: str,
                              content_catalog: Optional[Any],
                              poste_labels: Dict[str, str]) -> Optional[Any]:
        """
        Résout le contenu d'un poste en essayant code/label et variantes normalisées.
        """
        if not content_catalog:
            return None

        candidates = []
        if poste_code:
            candidates.append(poste_code)
        if poste_label and poste_label != poste_code:
            candidates.append(poste_label)

        reverse_labels = {label: code for code, label in poste_labels.items()}
        if poste_code in reverse_labels:
            candidates.append(reverse_labels[poste_code])
        if poste_label in reverse_labels:
            candidates.append(reverse_labels[poste_label])

        normalized = {c.strip().lower() for c in candidates if c}
        for key in candidates:
            content = content_catalog.get_content(key, activity)
            if content and content.poste_l1_code and content.poste_l1_code.strip().lower() in normalized:
                return content

        return None

    def _find_block_in_range(self, start_marker: str, end_marker: str,
                             parent_start: int, parent_end: int) -> Optional[Tuple[int, int]]:
        """
        Trouve un bloc entre deux indices de paragraphes.

        Args:
            start_marker: Marqueur de début (ex: '[[START_POST]]')
            end_marker: Marqueur de fin (ex: '[[END_POST]]')
            parent_start: Index de début de la zone
            parent_end: Index de fin de la zone

        Returns:
            Tuple (start_idx, end_idx) ou None si non trouvé
        """
        paragraphs = list(self.doc.paragraphs)
        start_idx = None

        for i in range(parent_start, min(parent_end + 1, len(paragraphs))):
            text = paragraphs[i].text.strip()
            if start_marker in text and start_idx is None:
                start_idx = i
            elif end_marker in text and start_idx is not None:
                return (start_idx, i)

        return None

    def _find_marker_index(self, marker: str, start_idx: int,
                           end_idx: Optional[int] = None) -> Optional[int]:
        """
        Trouve l'index du premier paragraphe contenant un marqueur.

        Args:
            marker: Marqueur à chercher
            start_idx: Index de départ
            end_idx: Index de fin (optionnel)

        Returns:
            Index du paragraphe ou None
        """
        paragraphs = list(self.doc.paragraphs)
        stop = end_idx if end_idx is not None else len(paragraphs)
        for i in range(start_idx, min(stop, len(paragraphs))):
            if marker in paragraphs[i].text:
                return i
        return None

    def _find_table_with_placeholder(self, placeholder: str):
        """Retourne la première table contenant un placeholder."""
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            return table
        return None

    def _clear_placeholder_in_table(self, table, placeholder: str):
        """Supprime un placeholder dans une table existante."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, '')

    def _find_paragraph_in_range(self, placeholder: str, start_idx: int, end_idx: int):
        """
        Retourne le premier paragraphe contenant un placeholder dans une zone spécifique.

        Args:
            placeholder: Placeholder à chercher
            start_idx: Index de début de la zone
            end_idx: Index de fin de la zone

        Returns:
            Paragraphe trouvé ou None
        """
        paragraphs = list(self.doc.paragraphs)
        for i in range(start_idx, min(end_idx + 1, len(paragraphs))):
            if placeholder in paragraphs[i].text:
                return paragraphs[i]
        return None

    def _find_paragraph_with_placeholder(self, placeholder: str):
        """Retourne le premier paragraphe contenant un placeholder."""
        for paragraph in self.doc.paragraphs:
            if placeholder in paragraph.text:
                return paragraph
        return None

    def _insert_table_after_paragraph(self, paragraph, cols: int):
        """Insère un tableau juste après un paragraphe."""
        table = self.doc.add_table(rows=0, cols=cols)
        tbl_element = table._element
        parent = tbl_element.getparent()
        if parent is not None:
            parent.remove(tbl_element)
        paragraph._element.addnext(tbl_element)
        return table

    def _filter_emissions_l2(self, poste_code: str, activity: str,
                             context: Dict[str, Any],
                             parent_tree_id: Optional[str] = None):
        """
        Filtre EMISSIONS_L2 pour un poste, une activité et une zone d'arbre.

        Args:
            poste_code: Code du poste L1
            activity: Activité (EU ou AEP)
            context: Contexte global
            parent_tree_id: Node ID parent (ORG ou LOT) pour filtrer les ENT

        Returns:
            DataFrame filtré ou None
        """
        emissions_l2_df = context.get('emissions_l2_df')
        if emissions_l2_df is None or emissions_l2_df.empty:
            return None

        df = emissions_l2_df

        if 'poste_l1_code' in df.columns:
            poste_filter = poste_code
            poste_labels = context.get('poste_labels', {})
            reverse_labels = {label: code for code, label in poste_labels.items()}
            if poste_code in reverse_labels:
                poste_filter = reverse_labels[poste_code]
            df = df[df['poste_l1_code'] == poste_filter]
        if df.empty:
            return None

        if 'activity' in df.columns:
            df = df[df['activity'] == activity]
        else:
            tree = context.get('tree')
            if parent_tree_id and tree and 'node_id' in df.columns:
                ent_ids = tree.get_ent_ids_by_activity(parent_tree_id, activity)
                if ent_ids:
                    df = df[df['node_id'].isin(ent_ids)]
            elif 'node_id' in df.columns and activity:
                df = df[df['node_id'].astype(str).str.contains(f'_{activity}_')]

        if df.empty:
            return None

        return df

    def _insert_post_content(self, block_start: int, block_end: int,
                            poste_code: str, activity: str,
                            content: Optional[Any], context: Dict[str, Any],
                            parent_tree_id: Optional[str] = None):
        """
        Insère graphiques, tableaux et images pour un bloc POST.
        N'affiche que les éléments qui ont une valeur dans la table TEXTE_RAPPORT.

        Args:
            block_start: Index de début du bloc
            block_end: Index de fin du bloc
            poste_code: Code du poste L1
            activity: Activité (EU ou AEP)
            content: Contenu du poste depuis ContentCatalog
            context: Contexte global
            parent_tree_id: ID du nœud parent (LOT ou ORG)
        """
        if not content:
            return

        # 1. Graphique (POST_CHART_1) - N'afficher que si chart_key a une valeur
        if hasattr(content, 'chart_key') and content.chart_key and content.chart_key.strip():
            content_catalog = context.get('content_catalog')
            if not content_catalog or content_catalog.is_chart_supported(content.chart_key):
                chart_buffer = self._generate_post_chart(
                    poste_code, activity, content.chart_key, context, parent_tree_id=parent_tree_id
                )
                if chart_buffer:
                    self._insert_image_in_range('{{POST_CHART_1}}', chart_buffer, block_start, block_end, width=5.0)

        # 2. Tableau (POST_TABLE_1) - N'afficher que si table_key a une valeur
        if hasattr(content, 'table_key') and content.table_key and content.table_key.strip():
            content_catalog = context.get('content_catalog')
            if not content_catalog or content_catalog.is_table_supported(content.table_key):
                self._insert_post_table(poste_code, activity, content.table_key, context,
                                        parent_tree_id=parent_tree_id, block_start=block_start, block_end=block_end)

        # 3. Image statique (POST_IMAGE_1) - N'afficher que si image_key a une valeur
        if hasattr(content, 'image_key') and content.image_key and content.image_key.strip():
            content_catalog = context.get('content_catalog')
            if content_catalog and content_catalog.is_image_supported(content.image_key):
                self._insert_asset_image_in_range('{{POST_IMAGE_1}}', content.image_key, block_start, block_end, width=5.0)

    def _process_post_blocks(self, parent_start_idx: int, parent_end_idx: int,
                            entity_key: str, activity: str, context: Dict[str, Any],
                            parent_tree_id: Optional[str] = None):
        """
        Traite les blocs POST imbriqués dans un bloc ACTIVITY.

        Args:
            parent_start_idx: Index de début du bloc ACTIVITY parent
            parent_end_idx: Index de fin du bloc ACTIVITY parent
            entity_key: Clé de l'entité (ex: 'LOT_STEP1_EU', 'ORG_EU')
            activity: Activité (EU ou AEP)
            context: Contexte global
        """
        from .word_blocks import BlockProcessor

        # Récupérer les données
        lot_results = context.get('lot_results', {})
        result = lot_results.get(entity_key)

        if not result or not hasattr(result, 'top_postes') or not result.top_postes:
            # Pas de résultats ou pas de top postes, supprimer le bloc POST
            block_info = self._find_block_in_range(
                '[[START_POST]]', '[[END_POST]]', parent_start_idx, parent_end_idx
            )
            if block_info:
                start_idx, end_idx = block_info
                self._delete_block(start_idx, end_idx)
            return

        top_n = context.get('top_n', 4)
        top_postes = result.top_postes[:top_n]
        poste_labels = context.get('poste_labels', {})
        content_catalog = context.get('content_catalog')

        # Trouver le bloc POST dans la zone parent
        processor = BlockProcessor(self.doc)
        block_info = self._find_block_in_range(
            '[[START_POST]]', '[[END_POST]]', parent_start_idx, parent_end_idx
        )

        if not block_info:
            return

        start_idx, end_idx = block_info

        # Dupliquer le bloc pour chaque top poste (sauf le premier qui existe déjà)
        if len(top_postes) > 1:
            processor.duplicate_block(start_idx, end_idx, len(top_postes) - 1)

        refreshed_end = self._find_marker_index('[[END_ACTIVITY]]', parent_start_idx)
        if refreshed_end is not None:
            parent_end_idx = refreshed_end

        # Re-scanner pour obtenir les nouveaux indices
        all_post_blocks = self._find_all_post_blocks(parent_start_idx, parent_end_idx)

        # Remplir chaque bloc POST (itération inverse pour éviter les décalages d'index)
        max_idx = min(len(top_postes), len(all_post_blocks)) - 1

        # Calculer le total de l'entité pour les pourcentages
        entity_total_tco2e = result.total_tco2e if result else 0

        for i in range(max_idx, -1, -1):
            poste_code, tco2e = top_postes[i]
            block_start, block_end = all_post_blocks[i]

            # Construire les remplacements
            poste_label = poste_labels.get(poste_code, poste_code)
            content = self._resolve_post_content(
                poste_code, poste_label, activity, content_catalog, poste_labels
            )

            # Calculer le pourcentage par rapport au total de l'entité
            percentage = (tco2e / entity_total_tco2e * 100) if entity_total_tco2e > 0 else 0

            # Formater avec la virgule française
            emissions_text = f"{tco2e:,.1f}".replace(",", " ").replace(".", ",")
            percentage_text = f"{percentage:.1f}".replace(".", ",")

            replacements = {
                '{{POST_TITLE}}': poste_label,
                '{{POST_TEXT}}': content.text if content and hasattr(content, 'text') else '',
                '{{emissions_poste}}': emissions_text,
                '{{emissions_poste_pourcentage}}': percentage_text,
            }

            # Remplacer dans le bloc
            processor.replace_in_block(block_start, block_end, replacements)

            # Gérer les graphiques, tableaux, images
            self._insert_post_content(block_start, block_end, poste_code, activity,
                                     content, context, parent_tree_id=parent_tree_id)

    def _generate_post_chart(self, poste_code: str, activity: str,
                            chart_key: str, context: Dict[str, Any],
                            parent_tree_id: Optional[str] = None) -> Optional[BytesIO]:
        """
        Génère un graphique pour un poste spécifique.

        Args:
            poste_code: Code du poste L1
            activity: Activité (EU ou AEP)
            chart_key: Clé du graphique (ex: 'TRAVAUX_BREAKDOWN')
            context: Contexte global

        Returns:
            BytesIO contenant l'image ou None
        """
        # Récupérer les données L2 pour ce poste
        filtered = self._filter_emissions_l2(
            poste_code, activity, context, parent_tree_id=parent_tree_id
        )

        if filtered is None or filtered.empty:
            return None

        required_cols = {'poste_l2', 'tco2e'}
        if not required_cols.issubset(set(filtered.columns)):
            return None

        filtered = filtered.groupby('poste_l2', as_index=False)['tco2e'].sum()

        try:
            # Générer le graphique via ChartGenerator
            # Note: ChartGenerator doit avoir une méthode générique ou spécifique par chart_key
            if chart_key == 'TRAVAUX_BREAKDOWN' and hasattr(self.chart_gen, 'generate_travaux_breakdown'):
                return self.chart_gen.generate_travaux_breakdown(filtered)
            elif chart_key == 'FILE_EAU_BREAKDOWN' and hasattr(self.chart_gen, 'generate_file_eau_breakdown'):
                return self.chart_gen.generate_file_eau_breakdown(filtered)
            elif chart_key == 'EM_INDIRECTES_SPLIT' and hasattr(self.chart_gen, 'generate_em_indirectes_split'):
                return self.chart_gen.generate_em_indirectes_split(filtered)
            elif chart_key == 'REACTIF_BREAKDOWN' and hasattr(self.chart_gen, 'generate_reactif_breakdown'):
                return self.chart_gen.generate_reactif_breakdown(filtered)
        except Exception:
            return None

        return None

    def _insert_post_table(self, poste_code: str, activity: str,
                          table_key: str, context: Dict[str, Any],
                          parent_tree_id: Optional[str] = None,
                          block_start: Optional[int] = None,
                          block_end: Optional[int] = None):
        """
        Insère un tableau pour un poste spécifique.

        Args:
            poste_code: Code du poste L1
            activity: Activité (EU ou AEP)
            table_key: Clé du tableau (ex: 'EM_INDIRECTES_TABLE')
            context: Contexte global
            parent_tree_id: ID du nœud parent
            block_start: Index de début de la zone de recherche (optionnel)
            block_end: Index de fin de la zone de recherche (optionnel)
        """
        filtered = self._filter_emissions_l2(
            poste_code, activity, context, parent_tree_id=parent_tree_id
        )
        if filtered is None or filtered.empty:
            return

        required_cols = {'poste_l2', 'tco2e'}
        if not required_cols.issubset(set(filtered.columns)):
            return

        filtered = filtered.groupby('poste_l2', as_index=False)['tco2e'].sum()
        filtered = filtered.sort_values('tco2e', ascending=False)

        # Si block_start/end sont fournis, chercher dans la zone spécifique
        if block_start is not None and block_end is not None:
            paragraph = self._find_paragraph_in_range('{{POST_TABLE_1}}', block_start, block_end)
        else:
            paragraph = self._find_paragraph_with_placeholder('{{POST_TABLE_1}}')

        if paragraph is None:
            return

        # Créer le tableau après le paragraphe
        table = self._insert_table_after_paragraph(paragraph, cols=2)
        self._delete_paragraph(paragraph)

        try:
            self.table_gen.generate_table(table_key, filtered[['poste_l2', 'tco2e']], table)
        except Exception:
            return

    def _delete_block(self, start_idx: int, end_idx: int):
        """
        Supprime un bloc complet (y compris les marqueurs).

        Args:
            start_idx: Index du paragraphe START
            end_idx: Index du paragraphe END
        """
        paragraphs = list(self.doc.paragraphs)

        # Supprimer tous les paragraphes du bloc (y compris les marqueurs)
        # Supprimer de la fin vers le début pour éviter les problèmes d'indices
        for i in range(end_idx, start_idx - 1, -1):
            if i < len(paragraphs):
                p = paragraphs[i]._element
                p.getparent().remove(p)

    def _find_all_activity_blocks(self, parent_start: int, parent_end: int) -> List[Tuple[int, int]]:
        """
        Trouve tous les blocs ACTIVITY dans une zone donnée.

        Args:
            parent_start: Index de début de la zone
            parent_end: Index de fin de la zone

        Returns:
            Liste de tuples (start_idx, end_idx) pour chaque bloc ACTIVITY
        """
        blocks = []
        current_idx = parent_start
        paragraphs = list(self.doc.paragraphs)

        while current_idx < parent_end and current_idx < len(paragraphs):
            start_idx = None
            end_idx = None

            for i in range(current_idx, min(parent_end, len(paragraphs))):
                text = paragraphs[i].text.strip()

                if '[[START_ACTIVITY]]' in text and start_idx is None:
                    start_idx = i
                elif '[[END_ACTIVITY]]' in text and start_idx is not None:
                    end_idx = i
                    break

            if start_idx is not None and end_idx is not None:
                blocks.append((start_idx, end_idx))
                current_idx = end_idx + 1
            else:
                break

        return blocks

    def _process_activity_blocks(self, parent_start_idx: int, parent_end_idx: int,
                                 parent_node_id: str, context: Dict[str, Any],
                                 parent_tree_id: Optional[str] = None):
        """
        Traite les blocs ACTIVITY imbriqués dans un bloc LOT ou au niveau ORG.

        Args:
            parent_start_idx: Index de début du bloc parent (LOT ou ORG)
            parent_end_idx: Index de fin du bloc parent
            parent_node_id: ID du nœud parent (ex: 'STEP1' pour LOT, 'ORG' pour ORG)
            context: Contexte global
        """
        from .word_blocks import BlockProcessor

        tree = context.get('tree')
        lot_results = context.get('lot_results', {})
        has_lots = context.get('has_lots', False)

        # Déterminer les activités à traiter
        if has_lots and parent_node_id != 'ORG':
            # Cas LOT
            activities = tree.get_lot_activities(parent_node_id) if tree else set()
            key_prefix = f"LOT_{parent_node_id}_"
            parent_tree_id = parent_node_id
        else:
            # Cas ORG (sans LOT)
            activities = tree.get_org_activities() if tree else set()
            key_prefix = "ORG_"
            if parent_tree_id is None and tree:
                parent_tree_id = tree.get_org().node_id

        # Convertir en liste ordonnée
        activities_list = sorted(list(activities))

        if not activities_list:
            # Pas d'activités, supprimer le bloc ACTIVITY
            block_info = self._find_block_in_range(
                '[[START_ACTIVITY]]', '[[END_ACTIVITY]]', parent_start_idx, parent_end_idx
            )
            if block_info:
                start_idx, end_idx = block_info
                self._delete_block(start_idx, end_idx)
            return

        # Trouver le bloc ACTIVITY dans la zone parent
        processor = BlockProcessor(self.doc)
        block_info = self._find_block_in_range(
            '[[START_ACTIVITY]]', '[[END_ACTIVITY]]', parent_start_idx, parent_end_idx
        )

        if not block_info:
            return

        start_idx, end_idx = block_info

        # Dupliquer pour chaque activité (sauf la première)
        if len(activities_list) > 1:
            processor.duplicate_block(start_idx, end_idx, len(activities_list) - 1)

        # Mettre à jour la fin du bloc parent après duplication
        if has_lots and parent_node_id != 'ORG':
            refreshed_end = self._find_marker_index('[[END_LOT]]', parent_start_idx)
            if refreshed_end is not None:
                parent_end_idx = refreshed_end

        # Re-scanner pour obtenir les nouveaux blocs ACTIVITY
        all_activity_blocks = self._find_all_activity_blocks(parent_start_idx, parent_end_idx)

        # Remplir chaque bloc ACTIVITY (itération inverse pour éviter les décalages d'index)
        max_idx = min(len(activities_list), len(all_activity_blocks)) - 1
        for i in range(max_idx, -1, -1):
            activity = activities_list[i]
            block_start, block_end = all_activity_blocks[i]

            # Clé pour récupérer les résultats
            entity_key = f"{key_prefix}{activity}"
            result = lot_results.get(entity_key)

            if not result:
                continue

            # Remplacer {{ENT_ACTIVITY}} + liste top postes + KPI entité
            activity_label = "Eau potable" if activity == "AEP" else "Eaux usées"
            top_postes_list = self._format_entity_top_postes_list(result, context.get('poste_labels', {}))

            # Calculer les KPI au niveau entité (LOT×ACTIVITÉ)
            indicator_results = context.get('indicator_results', {})
            indicator_result = indicator_results.get(entity_key)

            kpi_m3_entity = None
            kpi_hab_entity = None

            if indicator_result:
                kpi_m3_entity = self.kpi_calc.calculate_kpi_m3_entity(result, indicator_result, activity)
                kpi_hab_entity = self.kpi_calc.calculate_kpi_hab_entity(result, indicator_result)

            # Formater les KPI pour affichage (avec virgule française)
            kpi_m3_text = f"{kpi_m3_entity:.2f}".replace(".", ",") if kpi_m3_entity is not None else "N/A"
            kpi_hab_text = f"{kpi_hab_entity:.2f}".replace(".", ",") if kpi_hab_entity is not None else "N/A"

            replacements = {
                '{{ENT_ACTIVITY}}': activity_label,
                '{{ENTITY_TOP_POSTES_LIST}}': top_postes_list,
                '{{kpi_m3_lot_act}}': kpi_m3_text,
                '{{kpi_hab_lot_act}}': kpi_hab_text
            }

            processor.replace_in_block(block_start, block_end, replacements)

            # Traiter les blocs POST imbriqués (top N postes)
            self._process_post_blocks(block_start, block_end, entity_key, activity, context,
                                      parent_tree_id=parent_tree_id)

            refreshed_end = self._find_marker_index('[[END_ACTIVITY]]', block_start)
            if refreshed_end is not None:
                block_end = refreshed_end

            # Traiter les blocs OTHER_POST imbriqués (postes non top)
            self._process_other_post_blocks(block_start, block_end, entity_key, activity, context,
                                            parent_tree_id=parent_tree_id)

            # Insérer les graphiques au niveau entité (LOT×ACTIVITY)
            self._insert_entity_charts(block_start, block_end, entity_key, context)

    def _insert_entity_charts(self, block_start: int, block_end: int,
                              entity_key: str, context: Dict[str, Any]):
        """
        Insère les graphiques au niveau entité (LOT×ACTIVITY).

        Args:
            block_start: Index de début du bloc ACTIVITY
            block_end: Index de fin du bloc ACTIVITY
            entity_key: Clé de l'entité (ex: 'LOT_STEP1_EU', 'ORG_EU')
            context: Contexte global
        """
        lot_results = context.get('lot_results', {})
        result = lot_results.get(entity_key)

        if not result:
            return

        poste_labels = context.get('poste_labels', {})

        # 1. Chart scope entity (camembert scopes pour cette entité)
        scope_chart = self.chart_gen.generate_scope_pie_entity(result)
        if scope_chart:
            self._insert_image('{{chart_pie_scope_entity_activity}}', scope_chart, width=5.0)

        # 2. Chart postes entity (camembert postes L1 pour cette entité)
        postes_chart = self.chart_gen.generate_postes_pie_entity(result, poste_labels=poste_labels)
        if postes_chart:
            self._insert_image('{{chart_pie_postes_entity_activity}}', postes_chart, width=5.0)

    def _find_all_other_post_blocks(self, parent_start: int, parent_end: int) -> List[Tuple[int, int]]:
        """
        Trouve tous les blocs OTHER_POST dans une zone donnée.

        Args:
            parent_start: Index de début de la zone
            parent_end: Index de fin de la zone

        Returns:
            Liste de tuples (start_idx, end_idx) pour chaque bloc OTHER_POST
        """
        blocks = []
        current_idx = parent_start
        paragraphs = list(self.doc.paragraphs)

        while current_idx < parent_end and current_idx < len(paragraphs):
            start_idx = None
            end_idx = None

            for i in range(current_idx, min(parent_end, len(paragraphs))):
                text = paragraphs[i].text.strip()

                if '[[START_OTHER_POST]]' in text and start_idx is None:
                    start_idx = i
                elif '[[END_OTHER_POST]]' in text and start_idx is not None:
                    end_idx = i
                    break

            if start_idx is not None and end_idx is not None:
                blocks.append((start_idx, end_idx))
                current_idx = end_idx + 1
            else:
                break

        return blocks

    def _process_other_post_blocks(self, parent_start_idx: int, parent_end_idx: int,
                                    entity_key: str, activity: str, context: Dict[str, Any],
                                    parent_tree_id: Optional[str] = None):
        """
        Traite les blocs OTHER_POST imbriqués dans un bloc ACTIVITY.
        Affiche les postes qui ne sont PAS dans le top N.

        Args:
            parent_start_idx: Index de début du bloc ACTIVITY parent
            parent_end_idx: Index de fin du bloc ACTIVITY parent
            entity_key: Clé de l'entité (ex: 'LOT_STEP1_EU', 'ORG_EU')
            activity: Activité (EU ou AEP)
            context: Contexte global
        """
        from .word_blocks import BlockProcessor

        # Récupérer les données
        lot_results = context.get('lot_results', {})
        result = lot_results.get(entity_key)

        if not result or not hasattr(result, 'other_postes') or not result.other_postes:
            # Pas de résultats ou pas d'autres postes, supprimer le bloc OTHER_POST
            block_info = self._find_block_in_range(
                '[[START_OTHER_POST]]', '[[END_OTHER_POST]]', parent_start_idx, parent_end_idx
            )
            if block_info:
                start_idx, end_idx = block_info
                self._delete_block(start_idx, end_idx)
            return

        other_postes = result.other_postes
        poste_labels = context.get('poste_labels', {})
        content_catalog = context.get('content_catalog')

        # Trouver le bloc OTHER_POST dans la zone parent
        processor = BlockProcessor(self.doc)
        block_info = self._find_block_in_range(
            '[[START_OTHER_POST]]', '[[END_OTHER_POST]]', parent_start_idx, parent_end_idx
        )

        if not block_info:
            return

        start_idx, end_idx = block_info

        # Dupliquer le bloc pour chaque other poste (sauf le premier qui existe déjà)
        if len(other_postes) > 1:
            processor.duplicate_block(start_idx, end_idx, len(other_postes) - 1)

        refreshed_end = self._find_marker_index('[[END_ACTIVITY]]', parent_start_idx)
        if refreshed_end is not None:
            parent_end_idx = refreshed_end

        # Re-scanner pour obtenir les nouveaux indices
        all_other_post_blocks = self._find_all_other_post_blocks(parent_start_idx, parent_end_idx)

        # Remplir chaque bloc OTHER_POST (itération inverse pour éviter les décalages d'index)
        max_idx = min(len(other_postes), len(all_other_post_blocks)) - 1

        # Calculer le total de l'entité pour les pourcentages
        entity_total_tco2e = result.total_tco2e if result else 0

        for i in range(max_idx, -1, -1):
            poste_code, tco2e = other_postes[i]
            block_start, block_end = all_other_post_blocks[i]

            # Construire les remplacements
            poste_label = poste_labels.get(poste_code, poste_code)
            content = self._resolve_post_content(
                poste_code, poste_label, activity, content_catalog, poste_labels
            )

            # Calculer le pourcentage par rapport au total de l'entité
            percentage = (tco2e / entity_total_tco2e * 100) if entity_total_tco2e > 0 else 0

            # Formater la valeur tCO2e avec espace pour les milliers
            formatted_tco2e = self.kpi_calc.format_number(tco2e)

            # Formater avec la virgule française
            emissions_text = f"{tco2e:,.1f}".replace(",", " ").replace(".", ",")
            percentage_text = f"{percentage:.1f}".replace(".", ",")

            replacements = {
                '{{OTHER_POST_TITLE}}': poste_label,
                '{{OTHER_POST_TCO2E}}': formatted_tco2e,
                '{{OTHER_POST_TEXT}}': content.text if content and hasattr(content, 'text') else '',
                '{{emissions_poste}}': emissions_text,
                '{{emissions_poste_pourcentage}}': percentage_text,
            }

            # Remplacer dans le bloc
            processor.replace_in_block(block_start, block_end, replacements)

    def _clean_all_markers(self):
        """Nettoie tous les marqueurs de blocs du document."""
        from .word_blocks import BlockProcessor

        processor = BlockProcessor(self.doc)

        # Nettoyer dans l'ordre inverse (imbrication)
        processor.remove_block_markers('[[START_OTHER_POST]]', '[[END_OTHER_POST]]')
        processor.remove_block_markers('[[START_POST]]', '[[END_POST]]')
        processor.remove_block_markers('[[START_ACTIVITY]]', '[[END_ACTIVITY]]')
        processor.remove_block_markers('[[START_LOT]]', '[[END_LOT]]')

    def _insert_org_charts(self, context: Dict[str, Any]):
        """
        Insère les graphiques au niveau ORG.

        Args:
            context: Contexte global
        """
        org_result = context.get('org_result')
        if not org_result:
            return

        poste_labels = context.get('poste_labels', {})
        tree = context.get('tree')

        # 1. Chart emissions scope (camembert scopes)
        charts = {
            '{{chart_emissions_scope_org}}': self.chart_gen.generate_scope_pie(org_result),
        }

        # 2. Chart emissions total (PIE des postes L1) - CORRIGÉ
        charts['{{chart_emissions_total_org}}'] = self.chart_gen.generate_total_emissions_pie(
            org_result, poste_labels=poste_labels
        )

        # 3. Chart contribution LOT (bar chart)
        lot_results = context.get('lot_results', {})
        if lot_results and tree and tree.has_lots():
            # Pour chaque LOT, sommer EU + AEP
            lot_totals = {}
            for lot in tree.get_lots():
                total = 0.0
                for activity in ['EU', 'AEP']:
                    key = f"LOT_{lot.node_id}_{activity}"
                    if key in lot_results:
                        total += lot_results[key].total_tco2e
                if total > 0:
                    lot_totals[lot.node_name] = total

            if lot_totals:
                lot_data = list(lot_totals.items())
                charts['{{chart_contrib_lot}}'] = self.chart_gen.generate_lot_contribution(lot_data)

        # 4. Chart électricité par LOT (PIE répartition par LOT)
        # Chercher le poste "électricité" et agréger par LOT
        elec_by_lot = {}
        if tree and tree.has_lots():
            for lot in tree.get_lots():
                lot_elec_total = 0.0

                for activity in ['EU', 'AEP']:
                    key = f"LOT_{lot.node_id}_{activity}"
                    result = lot_results.get(key)

                    if result:
                        # Chercher le poste électricité
                        for poste_code, tco2e in result.emissions_by_poste.items():
                            if 'ELEC' in poste_code.upper() or 'ELECTRICITE' in poste_labels.get(poste_code, '').upper():
                                lot_elec_total += tco2e

                if lot_elec_total > 0:
                    elec_by_lot[lot.node_name] = lot_elec_total

        if elec_by_lot:
            charts['{{chart_emissions_elec_org}}'] = self.chart_gen.generate_elec_emissions_by_lot(elec_by_lot)

        # 5. Top 3 inter-lot (grouped bar chart) ou par ENT si pas de LOT
        top3_postes = []
        if org_result.top_postes:
            top3_postes = org_result.top_postes[:3]
        elif org_result.emissions_by_poste:
            sorted_postes = sorted(org_result.emissions_by_poste.items(),
                                   key=lambda x: x[1], reverse=True)
            top3_postes = sorted_postes[:3]

        if tree and top3_postes:
            top3_by_group = {}

            for poste_code, _ in top3_postes:
                poste_label = poste_labels.get(poste_code, poste_code)
                top3_by_group[poste_label] = {}

                if tree.has_lots():
                    for lot in tree.get_lots():
                        lot_total_for_poste = 0.0
                        for activity in ['EU', 'AEP']:
                            key = f"LOT_{lot.node_id}_{activity}"
                            result = lot_results.get(key)
                            if result and poste_code in result.emissions_by_poste:
                                lot_total_for_poste += result.emissions_by_poste[poste_code]
                        if lot_total_for_poste > 0:
                            top3_by_group[poste_label][lot.node_name] = lot_total_for_poste
                else:
                    emissions_df = context.get('emissions_df')
                    if emissions_df is None or emissions_df.empty:
                        continue
                    reverse_labels = {label: code for code, label in poste_labels.items()}
                    poste_filter = poste_code
                    if poste_code in reverse_labels:
                        poste_filter = reverse_labels[poste_code]
                    if 'poste_l1_code' in emissions_df.columns:
                        if poste_filter not in emissions_df['poste_l1_code'].unique():
                            poste_filter = poste_code
                    for ent in tree.get_ents():
                        ent_total = emissions_df[
                            (emissions_df['node_id'] == ent.node_id) &
                            (emissions_df['poste_l1_code'] == poste_filter)
                        ]['tco2e'].sum()
                        if ent_total > 0:
                            top3_by_group[poste_label][ent.node_name] = ent_total

            charts['{{chart_batonnet_inter_lot_top3}}'] = self.chart_gen.generate_inter_lot_top3(top3_by_group)

        # Insérer les graphiques dans le document
        for placeholder, img_buffer in charts.items():
            if img_buffer:
                self._insert_image(placeholder, img_buffer)

    def _insert_image_in_range(self, placeholder: str, img_buffer: BytesIO,
                               start_idx: int, end_idx: int,
                               width: Optional[float] = 5.0, height: Optional[float] = None):
        """
        Insère une image à la place d'un placeholder dans une zone spécifique du document.

        Args:
            placeholder: Placeholder à remplacer
            img_buffer: Buffer contenant l'image
            start_idx: Index de début de la zone de recherche
            end_idx: Index de fin de la zone de recherche
            width: Largeur en inches
            height: Hauteur en inches
        """
        paragraphs = list(self.doc.paragraphs)

        for i in range(start_idx, min(end_idx + 1, len(paragraphs))):
            paragraph = paragraphs[i]
            if placeholder in paragraph.text:
                # Supprimer le placeholder
                paragraph.clear()
                # Insérer l'image
                run = paragraph.add_run()
                run.add_picture(
                    img_buffer,
                    width=Inches(width) if width is not None else None,
                    height=Inches(height) if height is not None else None
                )
                return

    def _insert_asset_image_in_range(self, placeholder: str, image_key: str,
                                      start_idx: int, end_idx: int,
                                      width: Optional[float] = 5.0, height: Optional[float] = None):
        """
        Insère une image depuis assets/ dans une zone spécifique du document.

        Args:
            placeholder: Placeholder à remplacer (ex: '{{POST_IMAGE_1}}')
            image_key: Nom du fichier image (ex: 'DIGESTEUR_SCHEMA')
            start_idx: Index de début de la zone de recherche
            end_idx: Index de fin de la zone de recherche
            width: Largeur en inches
            height: Hauteur en inches
        """
        image_path = self.assets_path / f"{image_key}.png"

        if not image_path.exists():
            # Image non trouvée, laisser le placeholder (sera nettoyé)
            return

        # Charger l'image
        with open(image_path, 'rb') as f:
            img_bytes = f.read()

        img_buffer = BytesIO(img_bytes)

        # Insérer l'image dans la zone spécifiée
        self._insert_image_in_range(placeholder, img_buffer, start_idx, end_idx, width=width, height=height)

    def _insert_image(self, placeholder: str, img_buffer: BytesIO,
                      width: Optional[float] = 5.0, height: Optional[float] = None):
        """
        Insère une image à la place d'un placeholder (recherche globale dans tout le document).

        Args:
            placeholder: Placeholder à remplacer
            img_buffer: Buffer contenant l'image
            width: Largeur en inches
        """
        for paragraph in self.doc.paragraphs:
            if placeholder in paragraph.text:
                # Supprimer le placeholder
                paragraph.clear()
                # Insérer l'image
                run = paragraph.add_run()
                run.add_picture(
                    img_buffer,
                    width=Inches(width) if width is not None else None,
                    height=Inches(height) if height is not None else None
                )
                return

        # Chercher dans les tableaux
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            paragraph.clear()
                            run = paragraph.add_run()
                            run.add_picture(
                                img_buffer,
                                width=Inches(width) if width is not None else None,
                                height=Inches(height) if height is not None else None
                            )
                            return

    def _clean_empty_placeholders(self):
        """Supprime les paragraphes contenant des placeholders non remplacés."""
        paragraphs_to_remove = []

        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            # Si le paragraphe contient uniquement un placeholder {{...}}
            if re.match(r'^\{\{[A-Z_0-9]+\}\}$', text):
                paragraphs_to_remove.append(paragraph)

        # Supprimer les paragraphes
        for paragraph in paragraphs_to_remove:
            self._delete_paragraph(paragraph)

    def _delete_paragraph(self, paragraph):
        """Supprime un paragraphe du document."""
        p = paragraph._element
        p.getparent().remove(p)

    def save(self, output_path: str):
        """
        Sauvegarde le document rendu.

        Args:
            output_path: Chemin de sortie
        """
        if self.doc is None:
            raise ValueError("Document non chargé. Appelez render() d'abord.")

        self.doc.save(output_path)
